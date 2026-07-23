import io
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from lib.db import connection as db

router = APIRouter()

class ExportRequest(BaseModel):
    contract_id: str
    org_id: str
    include_suggestions: bool = True
    include_executive_summary: bool = True


# ---------------------------------------------------------------------------
# Theme — single source of truth for colors/fonts so nothing is hardcoded
# per-run. Change a value here and the whole document updates consistently.
# ---------------------------------------------------------------------------
class Theme:
    FONT = "Calibri"
    NAVY = "1B263B"
    MUTED = "6B7280"
    LIGHT_BG = "F7F9FC"
    WHITE = "FFFFFF"
    BORDER = "E2E8F0"

    RISK = {
        "critical": {"bg": "FDF2F2", "border": "E02424", "text": "9B1C1C", "label": "CRITICAL RISK"},
        "high":     {"bg": "FFF8F2", "border": "F79009", "text": "B45309", "label": "HIGH RISK"},
        "medium":   {"bg": "FAFAFA", "border": "CA8A04", "text": "854D0E", "label": "MEDIUM RISK"},
        "low":      {"bg": "FAFAFA", "border": "16A34A", "text": "166534", "label": "LOW RISK"},
    }


def render_clauses_section(doc, clauses, Theme, helpers):
    """
    Renders clauses as individually bordered/shaded cards instead of a
    continuous text flow. Each card = one table cell containing:
    header (clause # + risk pill) + clause text + explanation (if flagged).
    Using a single cell per clause means the left color bar renders as one
    continuous line regardless of how many paragraphs are inside it.
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    shade_cell = helpers["shade_cell"]
    left_border_cell = helpers["left_border_cell"]
    no_table_borders = helpers["no_table_borders"]
    set_cell_margins = helpers["set_cell_margins"]
    safe_text = helpers["safe_text"]

    NEUTRAL = {"bg": "FAFAFA", "border": "CBD5E1", "text": "64748B", "label": None}
    RISK_CFG = {**Theme.RISK, None: NEUTRAL, "unclassified": NEUTRAL}

    doc.add_paragraph("Analyzed Contract Clauses", style="LLSectionHeading")

    current_page = None
    clause_number = 0

    for c in clauses:
        clause_number += 1
        page_no = c.get("page_no", 1)

        if current_page != page_no:
            current_page = page_no
            marker = doc.add_paragraph(style="LLPageMarker")
            marker.paragraph_format.space_before = Pt(16)
            marker.add_run(f"SOURCE PAGE {current_page}")

        risk_label = c.get("risk_label")
        cfg = RISK_CFG.get(risk_label, NEUTRAL)
        show_explanation = risk_label in ("critical", "high")

        # --- Card container: single cell so the left border is unbroken ---
        card = doc.add_table(rows=1, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        card.autofit = False
        cell = card.rows[0].cells[0]
        cell.width = Inches(6.5)
        shade_cell(cell, cfg["bg"] if show_explanation else "FFFFFF")
        left_border_cell(cell, cfg["border"], size_pt=3.5 if show_explanation else 2)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200) # Maintain professional spacing

        # --- Card header: clause number (left) + risk pill (right-ish) ---
        header_p = cell.paragraphs[0]
        header_p.paragraph_format.space_before = Pt(8)
        header_p.paragraph_format.space_after = Pt(6)

        num_run = header_p.add_run(f"CLAUSE {clause_number:02d}")
        num_run.font.name = Theme.FONT
        num_run.font.size = Pt(8.5)
        num_run.font.bold = True
        num_run.font.color.rgb = RGBColor.from_string(Theme.MUTED)

        if cfg["label"]:
            sep_run = header_p.add_run("   ●   ")
            sep_run.font.size = Pt(8.5)
            sep_run.font.color.rgb = RGBColor.from_string(cfg["border"])

            pill_run = header_p.add_run(cfg["label"])
            pill_run.font.name = Theme.FONT
            pill_run.font.size = Pt(8.5)
            pill_run.font.bold = True
            pill_run.font.color.rgb = RGBColor.from_string(cfg["text"])

        # --- Clause body text ---
        body_p = cell.add_paragraph()
        body_p.paragraph_format.space_after = Pt(6 if show_explanation else 8)
        body_p.paragraph_format.line_spacing = 1.15
        body_run = body_p.add_run(safe_text(c.get("text"), "[No clause text extracted]"))
        body_run.font.name = Theme.FONT
        body_run.font.size = Pt(10.5)
        body_run.font.color.rgb = RGBColor.from_string("1A1A1A")

        # --- Explanation, only for flagged clauses, visually nested ---
        if show_explanation:
            divider_p = cell.add_paragraph()
            divider_p.paragraph_format.space_before = Pt(2)
            divider_p.paragraph_format.space_after = Pt(2)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            divider_p._p.get_or_add_pPr().append(parse_xml(
                f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" '
                f'w:space="1" w:color="{cfg["border"]}"/></w:pBdr>'
            ))

            why_p = cell.add_paragraph()
            why_p.paragraph_format.space_before = Pt(4)
            why_p.paragraph_format.space_after = Pt(8)
            why_label = why_p.add_run("WHY THIS MATTERS   ")
            why_label.font.name = Theme.FONT
            why_label.font.size = Pt(8)
            why_label.font.bold = True
            why_label.font.color.rgb = RGBColor.from_string(cfg["text"])

            why_text = why_p.add_run(safe_text(c.get("explanation"), "No explanation generated."))
            why_text.font.name = Theme.FONT
            why_text.font.size = Pt(9.5)
            why_text.font.italic = True
            why_text.font.color.rgb = RGBColor.from_string("464646")

        # Space between cards
        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_redline_docx(contract: dict, clauses: list[dict], req: ExportRequest) -> io.BytesIO:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    import json

    doc = Document()

    # -- Base styles (defined once, reused everywhere instead of per-run) --
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = Theme.FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    def add_style(name, size, bold=False, italic=False, color=Theme.NAVY, base="Normal", space_after=6):
        try:
            st = styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH
        except ValueError:
            st = styles[name]
        st.base_style = styles[base]
        st.font.name = Theme.FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = italic
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_after = Pt(space_after)
        return st

    add_style("LLTitle", 24, bold=True, space_after=4)
    add_style("LLSubtitle", 10, italic=True, color=Theme.MUTED, space_after=2)
    add_style("LLSectionHeading", 15, bold=True, space_after=10)
    add_style("LLPageMarker", 9, bold=True, color=Theme.MUTED, space_after=4)
    add_style("LLClauseBody", 11, space_after=4)
    add_style("LLCalloutHeader", 9.5, bold=True, space_after=2)
    add_style("LLCalloutBody", 9.5, italic=True, color="464646", space_after=6)

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -- Reusable helpers --
    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    def left_border_cell(cell, color_hex, size_pt=3.5):
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="nil"/>
                <w:left w:val="single" w:sz="{int(size_pt * 8)}" w:space="0" w:color="{color_hex}"/>
                <w:bottom w:val="nil"/>
                <w:right w:val="nil"/>
            </w:tcBorders>
        '''))

    def no_table_borders(table):
        table.style = None  # Detach default styles that carry cell borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")   # "nil" is more reliable than "none" across Word/LibreOffice
            borders.append(el)
        tblPr.append(borders)
        
        # Explicitly strip borders from all individual cells as well
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'''
                    <w:tcBorders {nsdecls("w")}>
                        <w:top w:val="nil"/>
                        <w:left w:val="nil"/>
                        <w:bottom w:val="nil"/>
                        <w:right w:val="nil"/>
                    </w:tcBorders>
                ''')
                tcPr.append(tcBorders)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(val))
            node.set(qn("w:type"), "dxa")
            mar.append(node)
        tcPr.append(mar)

    def add_page_number_field(paragraph):
        run = paragraph.add_run()
        for xml in (
            '<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
            '<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">PAGE</w:instrText>',
            '<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        ):
            run._r.append(parse_xml(xml))

    def safe_text(value, fallback="—"):
        return value if value not in (None, "") else fallback

    # -----------------------------------------------------------------
    # Header / Footer — consistent on every page instead of repeating
    # meta info inline in the body
    # -----------------------------------------------------------------
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hrun = hp.add_run(f"LegalLens  ·  Contract {req.contract_id}")
    hrun.font.name = Theme.FONT
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor.from_string(Theme.MUTED)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Page ")
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor.from_string(Theme.MUTED)
    add_page_number_field(fp)

    # -----------------------------------------------------------------
    # 1. Title block
    # -----------------------------------------------------------------
    title = doc.add_paragraph(style="LLTitle")
    title.add_run("Contract Review & Redline")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(style="LLSubtitle")
    subtitle.add_run(f"Contract ID: {req.contract_id}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # thin rule under the title instead of a spacer paragraph
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(16)
    pPr = rule_p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{Theme.BORDER}"/></w:pBdr>'
    ))

    # -----------------------------------------------------------------
    # 2. Executive summary with a visual risk bar (not just a table)
    # -----------------------------------------------------------------
    if req.include_executive_summary and contract:
        doc.add_paragraph("Executive Risk Summary", style="LLSectionHeading")

        try:
            breakdown = json.loads(contract.get("risk_breakdown") or "{}")
        except (TypeError, ValueError):
            breakdown = {}

        risk_score = contract.get("risk_score", 0)
        total_clauses = contract.get("total_clauses", len(clauses))
        flagged = contract.get("flagged_clauses", 0)

        counts = {k: int(breakdown.get(k, 0)) for k in ("critical", "high", "medium", "low")}
        total_flagged_bar = sum(counts.values()) or 1

        # Score line
        score_p = doc.add_paragraph()
        score_p.paragraph_format.space_after = Pt(2)
        r = score_p.add_run(f"{risk_score}")
        r.font.size = Pt(28)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(Theme.NAVY)
        r2 = score_p.add_run(" / 100 risk score")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(Theme.MUTED)

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(14)  # Give breathing room above the bar directly
        mr = meta_p.add_run(f"{total_clauses} clauses reviewed  ·  {flagged} flagged for attention")
        mr.font.size = Pt(9.5)
        mr.font.color.rgb = RGBColor.from_string(Theme.MUTED)

        # Proportional risk bar — one-row table with colored cells sized
        # by clause count, giving an at-a-glance distribution
        bar = doc.add_table(rows=1, cols=4)
        bar.alignment = WD_TABLE_ALIGNMENT.CENTER
        bar.autofit = False
        no_table_borders(bar)
        
        for i, key in enumerate(("critical", "high", "medium", "low")):
            frac = max(counts[key] / total_flagged_bar, 0.02) if counts[key] else 0
            width = Inches(6.5 * frac) if frac else Cm(0.05)
            cell = bar.rows[0].cells[i]
            cell.width = width
            bar.columns[i].width = width
            shade_cell(cell, Theme.RISK[key]["border"] if counts[key] else Theme.BORDER)
            
            # Apply thin white left/right borders to make segments clearly distinguishable
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="nil"/>
                    <w:left w:val="single" w:sz="12" w:space="0" w:color="FFFFFF"/>
                    <w:bottom w:val="nil"/>
                    <w:right w:val="single" w:sz="12" w:space="0" w:color="FFFFFF"/>
                </w:tcBorders>
            '''))
            
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.text = ""

        # Legend under the bar
        legend = doc.add_table(rows=1, cols=4)
        legend.alignment = WD_TABLE_ALIGNMENT.CENTER
        no_table_borders(legend)
        for i, key in enumerate(("critical", "high", "medium", "low")):
            cell = legend.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(12)  # Adjust spacing above the legend elements directly
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"● {key.capitalize()}  {counts[key]}")
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(Theme.RISK[key]["border"])

        # Page break directly after the legend (without loose trailing space)
        p_break = doc.add_paragraph()
        p_break.paragraph_format.space_before = Pt(0)
        p_break.paragraph_format.space_after = Pt(0)
        p_break.add_run().add_break()

    # -----------------------------------------------------------------
    # 3. Clauses
    # -----------------------------------------------------------------
    helpers = {
        "shade_cell": shade_cell,
        "left_border_cell": left_border_cell,
        "no_table_borders": no_table_borders,
        "set_cell_margins": set_cell_margins,
        "safe_text": safe_text,
    }
    render_clauses_section(doc, clauses, Theme, helpers)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


@router.post("/export/redline")
async def export_redline(req: ExportRequest):
    try:
        import docx  # noqa: F401 — validates dependency is installed before DB work
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    async with db.get_pool().acquire() as conn:
        contract = await conn.fetchrow(
            "SELECT * FROM contracts WHERE id = $1 AND org_id = $2",
            req.contract_id, req.org_id,
        )
        if not contract:
            raise HTTPException(status_code=404, detail="CONTRACT_NOT_FOUND")
            
        clauses = await conn.fetch(
            "SELECT * FROM clauses WHERE contract_id = $1 "
            "ORDER BY page_no, start_char ASC",
            req.contract_id,
        )

    if not clauses:
        raise HTTPException(status_code=422, detail="NO_CLAUSES_FOUND")

    bio = build_redline_docx(dict(contract), [dict(c) for c in clauses], req)

    import os
    file_name = contract.get("file_name", req.contract_id)
    base_name, _ = os.path.splitext(file_name)
    export_filename = f"redline_{base_name}.docx"

    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{export_filename}"'},
    )


@router.post("/export/summary")
async def export_summary(req: ExportRequest):
    import os
    import sys

    # Robust Windows DLL injection for GTK3/WeasyPrint
    if sys.platform == "win32":
        gtk_paths = [
            r"C:\Program Files\GTK3-Runtime\bin",
            r"C:\Program Files (x86)\GTK3-Runtime\bin",
        ]
        for p in gtk_paths:
            if os.path.isdir(p):
                try:
                    os.add_dll_directory(p)
                except Exception:
                    pass

    try:
        from weasyprint import HTML
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=(
                f"WeasyPrint is not fully set up on this system: {e}. "
                "Prerequisites: "
                "1. For Linux/Docker: Install libpango-1.0-0, libpangoft2-1.0-0, libjpeg-dev, and libffi-dev via apt. "
                "2. For Windows: Install GTK3 runtime from https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases and add the 'bin' folder to your PATH."
            )
        )

    import json
    import os

    # Fetch data
    async with db.get_pool().acquire() as conn:
        contract = await conn.fetchrow(
            "SELECT * FROM contracts WHERE id = $1 AND org_id = $2",
            req.contract_id, req.org_id,
        )
        if not contract:
            raise HTTPException(status_code=404, detail="CONTRACT_NOT_FOUND")
            
        clauses = await conn.fetch(
            "SELECT * FROM clauses WHERE contract_id = $1 "
            "ORDER BY page_no, start_char ASC",
            req.contract_id,
        )

    risk_score = contract.get("risk_score", 0)
    total_clauses = contract.get("total_clauses", len(clauses))
    flagged = contract.get("flagged_clauses", 0)

    try:
        breakdown = json.loads(contract.get("risk_breakdown") or "{}")
    except Exception:
        breakdown = {}

    critical_count = breakdown.get("critical", 0)
    high_count = breakdown.get("high", 0)
    medium_count = breakdown.get("medium", 0)
    low_count = breakdown.get("low", 0)

    # Filter clauses to show in summary (critical and high)
    flagged_clauses = [c for c in clauses if c.get("risk_label") in ("critical", "high")]

    # Build HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {{
                size: letter;
                margin: 0.8in;
                @bottom-right {{
                    content: "Page " counter(page);
                    font-family: 'Calibri', sans-serif;
                    font-size: 8.5pt;
                    color: #6B7280;
                }}
                @bottom-left {{
                    content: "LegalLens Contract Summary";
                    font-family: 'Calibri', sans-serif;
                    font-size: 8.5pt;
                    color: #6B7280;
                }}
            }}
            body {{
                font-family: 'Calibri', sans-serif;
                color: #1A1A1A;
                line-height: 1.4;
                margin: 0;
            }}
            .header-container {{
                border-bottom: 2px solid #E2E8F0;
                padding-bottom: 15px;
                margin-bottom: 30px;
                text-align: center;
            }}
            .title {{
                font-size: 24pt;
                color: #1B263B;
                margin: 0 0 5px 0;
                font-weight: bold;
            }}
            .subtitle {{
                font-size: 10pt;
                color: #6B7280;
                font-style: italic;
                margin: 0;
            }}
            .section-heading {{
                font-size: 14pt;
                font-weight: bold;
                color: #1B263B;
                margin-top: 25px;
                margin-bottom: 15px;
                border-bottom: 1px solid #E2E8F0;
                padding-bottom: 5px;
            }}
            .score-container {{
                display: flex;
                align-items: center;
                margin-bottom: 25px;
                background-color: #F7F9FC;
                padding: 15px;
                border-radius: 6px;
            }}
            .score-badge {{
                font-size: 32pt;
                font-weight: bold;
                color: #1B263B;
                margin-right: 20px;
                border-right: 2px solid #E2E8F0;
                padding-right: 20px;
            }}
            .score-text {{
                font-size: 11pt;
                color: #4A5568;
            }}
            .breakdown-table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 30px;
            }}
            .breakdown-table th, .breakdown-table td {{
                padding: 10px;
                text-align: left;
                font-size: 10pt;
            }}
            .breakdown-table th {{
                background-color: #1B263B;
                color: white;
                font-weight: bold;
            }}
            .breakdown-table tr:nth-child(even) td {{
                background-color: #F7F9FC;
            }}
            .card {{
                border-left: 4px solid #CBD5E1;
                background-color: #FAFAFA;
                padding: 12px 15px;
                margin-bottom: 15px;
                page-break-inside: avoid;
            }}
            .card.critical {{
                border-left-color: #E02424;
                background-color: #FDF2F2;
            }}
            .card.high {{
                border-left-color: #F79009;
                background-color: #FFF8F2;
            }}
            .card-header {{
                font-size: 8.5pt;
                font-weight: bold;
                color: #6B7280;
                margin-bottom: 6px;
            }}
            .card-header .critical-label {{ color: #9B1C1C; }}
            .card-header .high-label {{ color: #B45309; }}
            .card-body {{
                font-size: 10pt;
                margin-bottom: 8px;
            }}
            .divider {{
                border-top: 1px solid #E2E8F0;
                margin: 8px 0;
            }}
            .card-explanation {{
                font-size: 9.5pt;
                font-style: italic;
                color: #4A5568;
            }}
            .explanation-title {{
                font-size: 8pt;
                font-weight: bold;
                display: inline-block;
                margin-right: 5px;
            }}
            .critical-label {{ color: #E02424; }}
            .high-label {{ color: #F79009; }}
        </style>
    </head>
    <body>
        <div class="header-container">
            <h1 class="title">Contract Review Summary</h1>
            <p class="subtitle">Contract: {contract.get("file_name", req.contract_id)} (ID: {req.contract_id})</p>
        </div>

        <h2 class="section-heading">Executive Risk Assessment</h2>
        <div class="score-container">
            <div class="score-badge">{risk_score}</div>
            <div class="score-text">
                <strong>Overall Contract Risk Score</strong><br>
                This contract contains a total of {total_clauses} clauses, of which {flagged} clauses have been flagged for special review and potential risk.
            </div>
        </div>

        <table class="breakdown-table">
            <thead>
                <tr>
                    <th>Metric</th>
                    <th>Value</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td><strong>Risk Classification Breakdown</strong></td>
                    <td>
                        <span class="critical-label">● Critical: {critical_count}</span> &nbsp;|&nbsp;
                        <span class="high-label">● High: {high_count}</span> &nbsp;|&nbsp;
                        <span style="color: #CA8A04;">● Medium: {medium_count}</span> &nbsp;|&nbsp;
                        <span style="color: #16A34A;">● Low: {low_count}</span>
                    </td>
                </tr>
                <tr>
                    <td><strong>Total Clauses Reviewed</strong></td>
                    <td>{total_clauses}</td>
                </tr>
                <tr>
                    <td><strong>Flagged Risks (Critical & High)</strong></td>
                    <td>{critical_count + high_count}</td>
                </tr>
            </tbody>
        </table>

        <h2 class="section-heading">Key Flagged Risks</h2>
    """

    if not flagged_clauses:
        html_content += "<p>No critical or high-risk clauses detected in this agreement.</p>"
    else:
        for idx, c in enumerate(flagged_clauses):
            label = c.get("risk_label")
            class_name = label if label in ("critical", "high") else ""
            label_text = "CRITICAL RISK" if label == "critical" else "HIGH RISK"
            label_class = "critical-label" if label == "critical" else "high-label"

            html_content += f"""
            <div class="card {class_name}">
                <div class="card-header">
                    CLAUSE {idx+1:02d} &nbsp;●&nbsp; PAGE {c.get("page_no", 1)} &nbsp;●&nbsp; <span class="{label_class}">{label_text}</span>
                </div>
                <div class="card-body">
                    {c.get("text")}
                </div>
                <div class="divider"></div>
                <div class="card-explanation">
                    <span class="explanation-title {label_class}">WHY THIS MATTERS:</span>
                    {c.get("explanation", "No explanation generated.")}
                </div>
            </div>
            """

    html_content += """
    </body>
    </html>
    """

    # Compile PDF using WeasyPrint
    pdf_bytes = HTML(string=html_content).write_pdf()
    bio = io.BytesIO(pdf_bytes)
    bio.seek(0)

    # Build filename
    file_name = contract.get("file_name", req.contract_id)
    base_name, _ = os.path.splitext(file_name)
    export_filename = f"summary_{base_name}.pdf"

    return StreamingResponse(
        bio,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{export_filename}"'},
    )
